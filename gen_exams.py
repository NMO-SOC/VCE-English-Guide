#!/usr/bin/env python3
"""Pre-generate randomised practice exams by filling the original Exam Template.docx
(the same template the desktop ExamGenerator.exe used)."""
import json, os, random, re, shutil
from docx import Document
from docx.shared import Cm

BUILD = os.path.dirname(os.path.abspath(__file__))
EG = os.path.join(BUILD, "examgen")
OUT = os.path.join(EG, "out")
TEMPLATE = os.path.join(EG, "Exam Template.docx")
B_IMG = os.path.join(EG, "b")
C_IMG = os.path.join(EG, "c")
N_EXAMS = 16

data = json.load(open(os.path.join(EG, "examgen.json"), encoding="utf-8"))
os.makedirs(OUT, exist_ok=True)

def iter_paras(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def fill_text(p, ph, text):
    if ph not in p.text: return False
    full = p.text.replace(ph, text)
    for r in list(p.runs):
        r.text = ""
    lines = full.split("\n")
    if p.runs:
        base = p.runs[0]
    else:
        base = p.add_run()
    base.text = lines[0]
    for ln in lines[1:]:
        base.add_break()
        base.add_run = None  # noqa - keep single run; use paragraph-level run
        r = p.add_run(ln)
    return True

def fill_image(p, ph, img_path, width_cm):
    if ph not in p.text: return False
    for r in list(p.runs):
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return True

def two(pool):
    a = random.choice(pool)
    b = random.choice(pool)
    tries = 0
    while b == a and len(pool) > 1 and tries < 10:
        b = random.choice(pool); tries += 1
    return a, b

c_pool = [c for c in data["c"] if c["imgs"]]
for n in range(1, N_EXAMS + 1):
    doc = Document(TEMPLATE)
    sb_i, sb_ii = two(data["sb"])
    re_i, re_ii = two(data["re"])
    b = random.choice(data["b"])
    c = random.choice(c_pool)
    repl = {"{A7}": sb_i, "{A8}": sb_ii, "{A10}": re_i, "{A11}": re_ii,
            "{G9}": b["title"] or "", "{G11}": b["s1"] or "", "{G15}": b["s3"] or ""}
    imgs = {}
    if b["img"]:
        imgs["{G13}"] = (os.path.join(B_IMG, b["img"]), 10.5)
    for ph, f in zip(("{N13}", "{N16}", "{N17}"), c["imgs"]):
        imgs[ph] = (os.path.join(C_IMG, f), 15.5)
    done = set()
    for p in iter_paras(doc):
        for ph, txt in repl.items():
            if ph not in done and fill_text(p, ph, txt): done.add(ph)
        for ph, (f, w) in imgs.items():
            if ph not in done and fill_image(p, ph, f, w): done.add(ph)
    # clear any unfilled image placeholders (e.g. C set with <3 images)
    for p in iter_paras(doc):
        for ph in ("{G13}", "{N13}", "{N16}", "{N17}", "{A8}"):
            if ph in p.text and ph not in done:
                fill_text(p, ph, "" if ph != "{A8}" else sb_ii); done.add(ph)
    doc.save(os.path.join(OUT, "practice-exam-%02d.docx" % n))
    missing = {"{A7}","{A8}","{A10}","{A11}","{G9}","{G11}","{G15}"} - done
    if missing: print("exam %02d missing:" % n, missing)
print("generated", N_EXAMS, "exams")
