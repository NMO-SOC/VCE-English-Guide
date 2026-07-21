#!/usr/bin/env python3
"""Create web-template.docx: placeholders normalised into single runs, and grey
slot images inserted at the four image positions so the browser can swap bytes."""
import os, re, json, zipfile
from docx import Document
from docx.shared import Cm
from PIL import Image

BUILD = os.path.dirname(os.path.abspath(__file__))
EG = os.path.join(BUILD, "examgen")
SRC_T = os.path.join(EG, "Exam Template.docx")
if not os.path.exists(SRC_T):
    SRC_T = "/tmp/eg.exe_extracted/Exam Template.docx"
OUT = os.path.join(EG, "web-template.docx")

# unique slot pngs (distinct byte sizes for later identification)
SLOTS = {"{G13}": ("slot_g13", 1050, 700, Cm(10.5)),
         "{N13}": ("slot_n13", 1400, 1944, Cm(15.0)),
         "{N16}": ("slot_n16", 1402, 1946, Cm(15.0)),
         "{N17}": ("slot_n17", 1404, 1948, Cm(15.0))}
slot_files = {}
for ph, (name, w, h, _) in SLOTS.items():
    p = os.path.join("/tmp", name + ".png")
    img = Image.new("RGB", (w, h), (238, 238, 238))
    img.putpixel((0, 0), (237, 237, 237))
    img.save(p)
    slot_files[ph] = p

doc = Document(SRC_T)

def iter_paras(d):
    for p in d.paragraphs: yield p
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: yield p

PH = re.compile(r"\{[A-Z]\d+\}")
for p in iter_paras(doc):
    txt = p.text
    if not PH.search(txt): continue
    m = PH.search(txt)
    ph = m.group(0)
    if ph in SLOTS:
        for r in list(p.runs): r.text = ""
        run = p.runs[0] if p.runs else p.add_run()
        _, w, h, width = SLOTS[ph]
        run.add_picture(slot_files[ph], width=width)
    else:
        # normalise: whole placeholder into the first run
        intact = any(ph in r.text for r in p.runs)
        if not intact:
            for r in list(p.runs): r.text = ""
            (p.runs[0] if p.runs else p.add_run()).text = txt

doc.save(OUT)

# map slot -> media filename by byte size
sizes = {os.path.getsize(slot_files[ph]): ph for ph in SLOTS}
z = zipfile.ZipFile(OUT)
mapping = {}
for n in z.namelist():
    if n.startswith("word/media"):
        s = z.getinfo(n).file_size
        if s in sizes:
            mapping[sizes[s]] = n
xml = z.read("word/document.xml").decode("utf-8")
left = sorted(set(PH.findall(xml)))
print("slot mapping:", mapping)
print("text placeholders intact in xml:", left)
json.dump({"slots": {k.strip("{}").lower(): v for k, v in mapping.items()},
           "px": {k.strip("{}").lower(): [v[1], v[2]] for k, v in SLOTS.items()}},
          open(os.path.join(EG, "slots.json"), "w"))
